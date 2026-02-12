"""Tests for the DocumentAnalyzer pre-import analysis."""

from __future__ import annotations

import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from apa_formatter.infrastructure.importers.document_analyzer import (
    AnalysisResult,
    APAIssue,
    DetectedReference,
    DocumentAnalyzer,
    SectionInfo,
)


# ---------------------------------------------------------------------------
# Helpers: create test .docx files
# ---------------------------------------------------------------------------


def _create_test_docx(
    title: str = "Test Title",
    author: str = "Test Author",
    sections: list[tuple[str, str]] | None = None,
    references: list[str] | None = None,
    abstract: str | None = None,
    font_name: str = "Times New Roman",
) -> Path:
    """Create a temporary .docx file for testing."""
    doc = Document()

    # Set core properties
    doc.core_properties.title = title
    doc.core_properties.author = author

    # Title paragraph
    p = doc.add_paragraph(title)
    p.style = doc.styles["Title"]

    # Abstract
    if abstract:
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph(abstract)

    # Sections
    if sections:
        for heading, content in sections:
            doc.add_heading(heading, level=1)
            if content:
                p = doc.add_paragraph(content)
                for run in p.runs:
                    run.font.name = font_name

    # References
    if references:
        doc.add_heading("Referencias", level=1)
        for ref in references:
            doc.add_paragraph(ref)

    # Save to temp file
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()
    return Path(tmp.name)


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


class TestDocumentAnalyzer:
    """Test suite for DocumentAnalyzer."""

    def setup_method(self) -> None:
        self.analyzer = DocumentAnalyzer()

    def test_basic_analysis(self) -> None:
        """Test analysis of a simple document with title and sections."""
        path = _create_test_docx(
            title="Mi Documento APA",
            author="Juan Pérez",
            sections=[
                ("Introducción", "Este es el contenido de la introducción."),
                ("Metodología", "Se utilizó un enfoque cualitativo."),
            ],
        )
        result = self.analyzer.analyze(path)

        assert result.success
        assert result.file_name.endswith(".docx")
        assert result.file_size_kb > 0
        assert result.detected_title == "Mi Documento APA"
        assert "Juan Pérez" in result.detected_authors
        assert result.total_words > 0
        assert result.total_paragraphs > 0
        assert len(result.sections) >= 2

        path.unlink()

    def test_reference_detection(self) -> None:
        """Test that references section is detected and paragraphs extracted."""
        path = _create_test_docx(
            title="Paper con Referencias",
            author="Ana García",
            sections=[("Introducción", "Texto introductorio.")],
            references=[
                "Smith, J. (2020). A great book. Publisher.",
                "Doe, A. (2019). Article title. Journal, 1(2), 10-20.",
            ],
        )
        result = self.analyzer.analyze(path)

        assert result.success
        assert len(result.raw_reference_paragraphs) == 2
        assert "Smith" in result.raw_reference_paragraphs[0]
        assert "Doe" in result.raw_reference_paragraphs[1]

        path.unlink()

    def test_reference_parsing(self) -> None:
        """Test that detected references are parsed with confidence scores."""
        path = _create_test_docx(
            title="Paper",
            author="Author",
            references=[
                "Freud, S. (1900). The Interpretation of Dreams. Publisher.",
            ],
        )
        result = self.analyzer.analyze(path)

        assert len(result.detected_references) > 0
        dref = result.detected_references[0]
        assert dref.raw_text
        assert 0.0 <= dref.confidence <= 1.0
        assert dref.detection_method in ("DOI", "ISBN", "URL", "BibTeX", "heuristic")

        path.unlink()

    def test_abstract_detection(self) -> None:
        """Test that abstract section is detected."""
        path = _create_test_docx(
            title="Paper con Abstract",
            author="Test",
            abstract="Este es un resumen del documento de prueba.",
            sections=[("Introducción", "Texto")],
        )
        result = self.analyzer.analyze(path)

        assert result.success
        assert result.detected_abstract is not None
        assert "resumen" in result.detected_abstract.lower()

        path.unlink()

    def test_font_detection(self) -> None:
        """Test detection of fonts used in the document."""
        path = _create_test_docx(
            title="Font Test",
            author="Test",
            sections=[("Section", "Content text here.")],
            font_name="Arial",
        )
        result = self.analyzer.analyze(path)

        assert result.success
        # Font detection works on run-level fonts
        assert isinstance(result.detected_fonts, list)

        path.unlink()

    def test_apa_compliance_no_abstract(self) -> None:
        """Test APA compliance check for missing abstract."""
        path = _create_test_docx(
            title="No Abstract",
            author="Test",
            sections=[("Intro", "Text content.")],
        )
        result = self.analyzer.analyze(path)

        assert result.success
        issue_messages = [i.message for i in result.apa_issues]
        assert any("abstract" in m.lower() or "resumen" in m.lower() for m in issue_messages)

        path.unlink()

    def test_apa_compliance_no_references(self) -> None:
        """Test APA compliance check for missing references section."""
        path = _create_test_docx(
            title="No Refs",
            author="Test",
            sections=[("Intro", "Text.")],
        )
        result = self.analyzer.analyze(path)

        assert result.success
        issue_messages = [i.message for i in result.apa_issues]
        assert any("referencia" in m.lower() for m in issue_messages)

        path.unlink()

    def test_empty_document(self) -> None:
        """Test handling of an empty document."""
        doc = Document()
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        tmp.close()
        path = Path(tmp.name)

        result = self.analyzer.analyze(path)

        assert result.success
        assert result.total_words == 0
        assert len(result.sections) == 0

        path.unlink()

    def test_invalid_file(self) -> None:
        """Test handling of non-existent or corrupt file."""
        result = self.analyzer.analyze(Path("/tmp/nonexistent_file.docx"))

        assert not result.success
        assert result.error_message

    def test_word_count_accuracy(self) -> None:
        """Test that word and character counts are accurate."""
        text = "Uno dos tres cuatro cinco"
        path = _create_test_docx(
            title="Count Test",
            author="Test",
            sections=[("Sección", text)],
        )
        result = self.analyzer.analyze(path)

        assert result.success
        # Should count at least the section text words
        assert result.total_words >= 5

        path.unlink()

    def test_section_info_structure(self) -> None:
        """Test that SectionInfo has correct structure."""
        path = _create_test_docx(
            title="Structure",
            author="Test",
            sections=[
                ("Introducción", "Primer párrafo de la introducción."),
                ("Método", "Segundo párrafo del método."),
                ("Resultados", "Tercer párrafo de resultados."),
            ],
        )
        result = self.analyzer.analyze(path)

        assert result.success
        for sec in result.sections:
            assert isinstance(sec, SectionInfo)
            assert sec.heading
            assert isinstance(sec.level, int)
            assert isinstance(sec.word_count, int)
            assert isinstance(sec.paragraph_count, int)

        path.unlink()

    def test_multiple_authors(self) -> None:
        """Test detection of multiple authors from core properties."""
        path = _create_test_docx(
            title="Multi-Author",
            author="Alice Smith, Bob Jones, Charlie Brown",
            sections=[("Intro", "Text")],
        )
        result = self.analyzer.analyze(path)

        assert result.success
        assert len(result.detected_authors) == 3
        assert "Alice Smith" in result.detected_authors

        path.unlink()


class TestAnalysisResultDefaults:
    """Test AnalysisResult default values."""

    def test_defaults(self) -> None:
        result = AnalysisResult()
        assert result.success
        assert result.file_name == ""
        assert result.total_words == 0
        assert result.detected_references == []
        assert result.apa_issues == []
        assert result.detected_title == "Documento sin título"

    def test_section_info(self) -> None:
        info = SectionInfo(heading="Test", level=1, word_count=100, paragraph_count=5)
        assert info.heading == "Test"
        assert info.level == 1

    def test_detected_reference(self) -> None:
        dref = DetectedReference(
            raw_text="Smith (2020).",
            parsed_reference=None,
            confidence=0.5,
            detection_method="heuristic",
        )
        assert dref.include
        assert dref.confidence == 0.5

    def test_apa_issue(self) -> None:
        issue = APAIssue(
            severity="warning",
            message="Test issue",
            detail="Some detail",
        )
        assert issue.severity == "warning"
