from pathlib import Path
from docx import Document
from apa_formatter.converters.extractor import extract_content_with_formatting
from apa_formatter.adapters.docx_adapter import DocxAdapter


def verify_roundtrip():
    # 1. Create a formatted DOCX
    src_path = Path("test_formatting.docx")
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Normal text. ")
    p.add_run("Bold text.").bold = True
    p.add_run(" ")
    p.add_run("Italic text.").italic = True
    doc.save(src_path)

    # 2. Extract
    print(f"Extracting from {src_path}...")
    apa_doc = extract_content_with_formatting(src_path)

    # Check extracted content
    section = apa_doc.sections[0]
    content = section.content
    print(f"Extracted content: {content!r}")

    assert "**Bold text.**" in content or "**Bold text**" in content
    assert "*Italic text.*" in content or "*Italic text*" in content

    # 3. Export
    out_path = Path("test_formatting_out.docx")
    print(f"Exporting to {out_path}...")
    adapter = DocxAdapter(apa_doc)
    adapter.generate(out_path)

    # 4. Verify Export
    doc_out = Document(out_path)
    found_bold = False
    found_italic = False

    for p in doc_out.paragraphs:
        for run in p.runs:
            if run.bold and "Bold text" in run.text:
                found_bold = True
            if run.italic and "Italic text" in run.text:
                found_italic = True

    print(f"Found bold: {found_bold}")
    print(f"Found italic: {found_italic}")

    assert found_bold, "Export failed to preserve bold"
    assert found_italic, "Export failed to preserve italic"

    print("✅ Verification passed due to successful round-trip!")

    # Cleanup
    if src_path.exists():
        src_path.unlink()
    if out_path.exists():
        out_path.unlink()


def verify_example_import():
    example_path = Path("Examples/Especificación Requisitos.docx")
    if not example_path.exists():
        print(f"Skipping example test: {example_path} not found")
        return

    print(f"\nTesting import of {example_path}...")
    try:
        doc = extract_content_with_formatting(example_path)
        print("✅ Successfully imported example")
        print(f"   Title: {doc.title_page.title}")
        print(f"   Authors: {doc.title_page.authors}")
        print(f"   Sections: {len(doc.sections)}")
        if doc.sections:
            print(f"   First Section: {doc.sections[0].heading}")
            print(f"   First Section Content Preview: {doc.sections[0].content[:100]}...")
    except Exception as e:
        print(f"❌ Failed to import example: {e}")
        raise e


if __name__ == "__main__":
    verify_roundtrip()
    verify_example_import()
