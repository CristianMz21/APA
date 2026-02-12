"""DOCX import dialog.

Allows the user to select an existing .docx file, extracts its content
using the converter module, and returns an APADocument that populates
the structured editor.
"""

from __future__ import annotations

from pathlib import Path

from PySide6.QtWidgets import (
    QDialog,
    QDialogButtonBox,
    QFileDialog,
    QHBoxLayout,
    QLabel,
    QPushButton,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)

from apa_formatter.models.document import APADocument


class ImportDialog(QDialog):
    """Dialog for importing an existing .docx file into the editor."""

    def __init__(self, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self.setWindowTitle("📥 Importar Documento .docx")
        self.setMinimumSize(500, 350)

        self._result_doc: APADocument | None = None

        layout = QVBoxLayout(self)

        # ── Info ──────────────────────────────────────────────────────────
        info = QLabel(
            "Seleccione un documento .docx existente para importar.\n"
            "El contenido se extraerá y cargará en el editor estructurado."
        )
        info.setWordWrap(True)
        info.setStyleSheet("font-size: 10pt; margin-bottom: 8px;")
        layout.addWidget(info)

        # ── File selector ─────────────────────────────────────────────────
        file_row = QHBoxLayout()
        self._path_label = QLabel("Ningún archivo seleccionado")
        self._path_label.setStyleSheet("color: #888; font-style: italic; padding: 4px;")
        btn_browse = QPushButton("📂 Seleccionar .docx")
        btn_browse.clicked.connect(self._on_browse)
        file_row.addWidget(self._path_label, stretch=1)
        file_row.addWidget(btn_browse)
        layout.addLayout(file_row)

        # ── Log ───────────────────────────────────────────────────────────
        layout.addWidget(QLabel("Resultado de la importación:"))
        self._log = QTextEdit()
        self._log.setReadOnly(True)
        self._log.setMaximumHeight(140)
        self._log.setStyleSheet(
            "font-family: monospace; font-size: 9pt;background: #FAFAFA; border: 1px solid #DDDDDD;"
        )
        layout.addWidget(self._log)

        # ── Buttons ───────────────────────────────────────────────────────
        self._buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        self._buttons.button(QDialogButtonBox.StandardButton.Ok).setEnabled(False)
        self._buttons.accepted.connect(self.accept)
        self._buttons.rejected.connect(self.reject)
        layout.addWidget(self._buttons)

        self.setStyleSheet(_IMPORT_STYLE)

    def _on_browse(self) -> None:
        path, _ = QFileDialog.getOpenFileName(
            self, "Abrir documento .docx", "", "Word Documents (*.docx)"
        )
        if not path:
            return
        self._path_label.setText(path)
        self._path_label.setStyleSheet("color: #333;")
        self._try_import(Path(path))

    def _try_import(self, path: Path) -> None:
        self._log.clear()
        try:
            from apa_formatter.converters.extractor import extract_content_with_formatting

            doc = extract_content_with_formatting(path)
            self._result_doc = doc
            self._log.append(f"✅ Importado con formato: {path.name}")
            self._log.append(f"   Título: {doc.title_page.title}")
            self._log.append(f"   Autores: {', '.join(doc.title_page.authors)}")
            self._log.append(f"   Secciones: {len(doc.sections)}")
            if doc.references:
                self._log.append(f"   Referencias: {len(doc.references)}")
            self._buttons.button(QDialogButtonBox.StandardButton.Ok).setEnabled(True)
        except ImportError:
            # Fallback: try basic extraction
            self._try_basic_import(path)
        except Exception as exc:
            self._log.append(f"❌ Error al importar: {exc}")
            self._result_doc = None
            self._buttons.button(QDialogButtonBox.StandardButton.Ok).setEnabled(False)

    def _try_basic_import(self, path: Path) -> None:
        """Fallback extraction using python-docx directly."""
        try:
            from docx import Document as DocxDocument

            from apa_formatter.models.document import (
                APADocument,
                Section,
                TitlePage,
            )
            from apa_formatter.models.enums import HeadingLevel

            docx_doc = DocxDocument(str(path))

            # Extract title from first paragraph
            title = "Documento Importado"
            sections = []
            current_content = []

            for para in docx_doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style_name = (para.style.name or "").lower()

                if "title" in style_name:
                    title = text
                elif "heading" in style_name:
                    # Save previous section
                    if current_content:
                        sections.append(
                            Section(
                                heading=sections[-1].heading if sections else "Introducción",
                                level=HeadingLevel.LEVEL_1,
                                content="\n".join(current_content),
                            )
                        )
                        current_content = []
                    sections.append(
                        Section(
                            heading=text,
                            level=HeadingLevel.LEVEL_1,
                            content="",
                        )
                    )
                else:
                    current_content.append(text)

            # Final section content
            if current_content and sections:
                last = sections[-1]
                sections[-1] = Section(
                    heading=last.heading,
                    level=last.level,
                    content=last.content + "\n".join(current_content),
                )

            if not sections and current_content:
                sections.append(
                    Section(
                        heading="Contenido",
                        level=HeadingLevel.LEVEL_1,
                        content="\n".join(current_content),
                    )
                )

            self._result_doc = APADocument(
                title_page=TitlePage(title=title, authors=["Autor"]),
                sections=sections,
            )
            self._log.append(f"✅ Importación básica: {path.name}")
            self._log.append(f"   Título: {title}")
            self._log.append(f"   Secciones: {len(sections)}")
            self._buttons.button(QDialogButtonBox.StandardButton.Ok).setEnabled(True)

        except Exception as exc:
            self._log.append(f"❌ Error en importación básica: {exc}")
            self._result_doc = None
            self._buttons.button(QDialogButtonBox.StandardButton.Ok).setEnabled(False)

    def get_document(self) -> APADocument | None:
        return self._result_doc


_IMPORT_STYLE = """
QDialog { background: white; }
QPushButton {
    background: #4A90D9;
    color: white;
    border: none;
    border-radius: 3px;
    padding: 6px 14px;
    font-size: 10pt;
}
QPushButton:hover { background: #357ABD; }
QPushButton:disabled { background: #CCCCCC; }
"""
