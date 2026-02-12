"""DOCX semantic parser strategy.

Reads a ``.docx`` file using ``python-docx`` and produces a list of
``ContentBlock`` objects — an intermediate representation that preserves
paragraph-level formatting metadata (alignment, bold, font, page index)
for downstream analysis handlers.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.exceptions import PackageNotFoundError

# ---------------------------------------------------------------------------
# Heading style regex (matches "Heading 1", "heading 2", etc.)
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(r"heading\s*(\d)", re.IGNORECASE)


# ---------------------------------------------------------------------------
# ContentBlock — the intermediate representation
# ---------------------------------------------------------------------------


@dataclass
class ContentBlock:
    """A single paragraph extracted from a DOCX with rich metadata.

    Each block carries enough formatting information for the analysis
    handlers to decide what role the paragraph plays in the document
    structure (title, heading, body, reference, etc.).
    """

    text: str
    style_name: str = ""
    alignment: int | None = None  # WD_ALIGN_PARAGRAPH value
    is_bold: bool = False
    is_italic: bool = False
    font_name: str | None = None
    font_size_pt: float | None = None
    page_index: int = 0
    heading_level: int | None = None
    is_list_item: bool = False
    has_page_break_before: bool = False
    raw_runs: list[dict[str, object]] = field(default_factory=list)
    # Table support
    is_table: bool = False
    table_data: list[list[str]] | None = None

    @property
    def is_centered(self) -> bool:
        """True if the paragraph has CENTER alignment."""
        return self.alignment == WD_ALIGN_PARAGRAPH.CENTER

    @property
    def is_heading(self) -> bool:
        """True if a heading level was detected."""
        return self.heading_level is not None

    @property
    def is_empty(self) -> bool:
        return not self.text.strip()


# ---------------------------------------------------------------------------
# DOCX Semantic Parser
# ---------------------------------------------------------------------------


class DocxSemanticParser:
    """Parses a ``.docx`` file into a list of ``ContentBlock`` objects.

    This is a *Strategy* — different file formats (PDF, plain text) can
    provide their own parser that produces the same ``ContentBlock`` list.
    """

    def parse(self, path: Path) -> list[ContentBlock]:
        """Read *path* and return enriched content blocks.

        Raises ``ValueError`` if the file cannot be opened.
        """
        try:
            doc = Document(str(path))
        except PackageNotFoundError:
            raise ValueError(f"No se pudo abrir el archivo: {path.name}")
        except Exception as exc:
            raise ValueError(f"Error al abrir el archivo: {exc}") from exc

        blocks: list[ContentBlock] = []
        current_page: int = 0

        # Helper to process a paragraph element
        def process_paragraph(para) -> ContentBlock | None:
            nonlocal current_page

            # Track page breaks
            has_break = self._has_page_break(para)
            if has_break:
                current_page += 1

            text = para.text.strip()

            # Style information
            style_name = (para.style.name if para.style else "").lower()
            alignment = para.alignment

            # Run-level aggregation
            runs_with_text = [r for r in para.runs if r.text.strip()]
            is_bold = bool(runs_with_text) and all(r.bold for r in runs_with_text)
            is_italic = bool(runs_with_text) and all(r.italic for r in runs_with_text)

            # Font info from first run
            font_name: str | None = None
            font_size_pt: float | None = None
            if runs_with_text:
                first_run = runs_with_text[0]
                if first_run.font:
                    font_name = first_run.font.name
                    if first_run.font.size:
                        font_size_pt = first_run.font.size.pt

            # Heading level
            heading_level = self._detect_heading_level(style_name, para)

            # List item?
            is_list = "list" in style_name

            # Raw runs for detailed inspection
            raw_runs = [
                {
                    "text": r.text,
                    "bold": r.bold,
                    "italic": r.italic,
                    "font_name": r.font.name if r.font else None,
                    "font_size": r.font.size.pt if r.font and r.font.size else None,
                }
                for r in para.runs
            ]

            return ContentBlock(
                text=text,
                style_name=style_name,
                alignment=alignment,
                is_bold=is_bold,
                is_italic=is_italic,
                font_name=font_name,
                font_size_pt=font_size_pt,
                page_index=current_page,
                heading_level=heading_level,
                is_list_item=is_list,
                has_page_break_before=has_break,
                raw_runs=raw_runs,
            )

        # Helper to process a table element
        def process_table(table) -> ContentBlock:
            rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                rows.append(row_data)

            # Linearize table text for simple semantic analysis
            # e.g. "Row 1: Col1 | Col2... \n Row 2..."
            text_lines = []
            for row in rows:
                text_lines.append(" | ".join(row))
            table_text = "\n".join(text_lines)

            return ContentBlock(
                text=table_text, page_index=current_page, is_table=True, table_data=rows
            )

        # Iterate over document body elements in order
        # doc.element.body is the w:body element
        # It contains w:p (paragraphs) and w:tbl (tables) as children
        from docx.text.paragraph import Paragraph
        from docx.table import Table

        if doc.element.body:
            for child in doc.element.body.iterchildren():
                if child.tag.endswith("p"):  # Paragraph
                    # We need to wrap the xml element in a Paragraph object
                    # But how? python-docx creates them internally.
                    #
                    # Better Approach:
                    # python-docx has iter_block_items() in newer versions?
                    # Or we can iterate doc.paragraphs and doc.tables and merge by element order?

                    # Reliable method: Use doc.iter_inner_content() if available,
                    # otherwise iterate elements and check parent. Note: python-docx doesn't expose convenient wrapping.

                    # Let's iterate doc.element.body and construct wrappers or match by identity
                    pass

        # Since we can't easily wrap XML elements back to python-docx objects without internals,
        # we'll use a pragmatic approach:
        # 1. Collect all paragraphs and tables with their XML element as key.
        # 2. Iterate XML body children.
        # 3. Lookup the corresponding python-docx object.

        para_map = {p._element: p for p in doc.paragraphs}
        table_map = {t._element: t for t in doc.tables}

        if doc.element.body:
            for child in doc.element.body.iterchildren():
                if child in para_map:
                    block = process_paragraph(para_map[child])
                    if block:
                        blocks.append(block)
                elif child in table_map:
                    block = process_table(table_map[child])
                    blocks.append(block)

        # Also extract page dimensions from sections
        self._page_dims: dict[str, float] | None = None
        try:
            section = doc.sections[0]
            self._page_dims = {
                "width_cm": section.page_width.cm if section.page_width else 21.59,
                "height_cm": section.page_height.cm if section.page_height else 27.94,
            }
        except Exception:
            pass

        # Font/spacing summary
        self._all_fonts: set[str] = set()
        self._spacings: list[float] = []
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font and run.font.name:
                    self._all_fonts.add(run.font.name)
            fmt = para.paragraph_format
            if fmt and fmt.line_spacing is not None:
                try:
                    sp = float(fmt.line_spacing)
                    if 0.5 < sp < 5.0:
                        self._spacings.append(sp)
                except (TypeError, ValueError):
                    pass

        return blocks

    @property
    def page_dimensions(self) -> dict[str, float] | None:
        """Page width/height in cm from the first DOCX section."""
        return self._page_dims

    @property
    def detected_fonts(self) -> list[str]:
        """All font names found across document runs."""
        return sorted(self._all_fonts) if self._all_fonts else []

    @property
    def dominant_line_spacing(self) -> float | None:
        """Most common line spacing value, or None."""
        if not self._spacings:
            return None
        from collections import Counter

        counter = Counter(self._spacings)
        return counter.most_common(1)[0][0]

    # -- Private helpers -----------------------------------------------------

    @staticmethod
    def _detect_heading_level(style_name: str, para) -> int | None:
        """Detect heading level from style name or bold heuristic."""
        match = _HEADING_RE.search(style_name)
        if match:
            return int(match.group(1))

        # Heuristic: short + all bold + Normal style → level 1
        if "normal" in style_name and len(para.text) < 100:
            runs_with_text = [r for r in para.runs if r.text.strip()]
            if runs_with_text and all(r.bold for r in runs_with_text):
                return 1

        return None

    @staticmethod
    def _has_page_break(para) -> bool:
        """Check if *para* contains a page break."""
        for run in para.runs:
            xml = getattr(run._element, "xml", "") or ""
            if "w:br" in xml and 'w:type="page"' in xml:
                return True
        # Also check paragraph-level page break before
        fmt = para.paragraph_format
        if fmt and fmt.page_break_before:
            return True
        return False
