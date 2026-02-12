"""DOCX content extractor with formatting support.

Extracts content from .docx files preserving structure, heading levels,
abstract detection, and basic formatting (bold → **text**, italic → *text*).
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from apa_formatter.models.document import APADocument, Section, TitlePage
from apa_formatter.models.enums import HeadingLevel


_ABSTRACT_HEADINGS = {"abstract", "resumen"}

_HEADING_RE = re.compile(r"heading\s*(\d)", re.IGNORECASE)


def extract_content_with_formatting(path: Path) -> APADocument:
    """Extract content from a .docx file, preserving basic formatting.

    Converts bold to **text** and italic to *text*.
    Detects real heading levels (H1, H2, H3) instead of flattening.
    Extracts abstract and keywords if found.
    """
    try:
        doc = Document(str(path))
    except PackageNotFoundError:
        raise ValueError(f"Could not open file: {path}")

    # 1. Metadata Extraction (Best Effort)
    title = "Untitled Document"
    authors = ["Unknown Author"]
    affiliation = "Unknown Affiliation"

    # Try core properties first
    try:
        props = doc.core_properties
        if props.author and props.author.strip():
            raw = props.author.strip()
            parsed_authors = re.split(r"[;,]|\band\b", raw)
            parsed_authors = [a.strip() for a in parsed_authors if a.strip()]
            if parsed_authors:
                authors = parsed_authors
        if props.title and props.title.strip():
            title = props.title.strip()
    except Exception:
        pass

    # Try to find title from paragraph styles
    title_from_style = False
    for para in doc.paragraphs[:5]:
        if "Title" in (para.style.name if para.style else ""):
            title = para.text.strip()
            title_from_style = True
            break

    # Fallback: first centered bold paragraph
    if not title_from_style and title == "Untitled Document":
        for para in doc.paragraphs[:10]:
            text = para.text.strip()
            if not text:
                continue
            if para.alignment == 1 and any(r.bold for r in para.runs):
                title = text
                break

    # 2. Section Extraction with real heading levels
    sections: list[Section] = []
    current_heading: str | None = None
    current_level: HeadingLevel = HeadingLevel.LEVEL_1
    current_content: list[str] = []

    abstract_text: str | None = None
    keywords: list[str] = []
    in_abstract = False

    for para in doc.paragraphs:
        style_name = ((para.style.name) if para.style else "").lower()
        text = _extract_formatted_text(para).strip()

        if not text:
            continue

        # Detect Headings
        level = _detect_heading_level(style_name, para)

        if level:
            # Check if this is an abstract heading
            clean_heading = text.replace("**", "").replace("*", "").strip()
            if clean_heading.lower() in _ABSTRACT_HEADINGS:
                in_abstract = True
                # Save previous section first
                if current_heading or current_content:
                    sections.append(
                        Section(
                            heading=current_heading,
                            level=current_level,
                            content="\n\n".join(current_content),
                        )
                    )
                current_heading = None
                current_content = []
                continue

            # End abstract collection if we hit a new heading
            if in_abstract:
                in_abstract = False

            # Save previous section
            if current_heading or current_content:
                sections.append(
                    Section(
                        heading=current_heading,
                        level=current_level,
                        content="\n\n".join(current_content),
                    )
                )

            # Start new section with detected level
            current_heading = clean_heading
            current_level = level
            current_content = []
        else:
            # Body text
            if in_abstract:
                # Check for keywords line
                kw_match = re.match(
                    r"(?:\*{0,2})(?:keywords?|palabras?\s*clave)(?:\*{0,2}):\s*(.+)",
                    text,
                    re.IGNORECASE,
                )
                if kw_match:
                    kw_text = kw_match.group(1).strip().rstrip(".")
                    keywords = [k.strip() for k in kw_text.split(",") if k.strip()]
                    in_abstract = False
                else:
                    if abstract_text:
                        abstract_text += " " + text.replace("**", "").replace("*", "")
                    else:
                        abstract_text = text.replace("**", "").replace("*", "")
                continue

            # Handle lists
            if "list" in style_name:
                text = f"- {text}"

            current_content.append(text)

    # Save last section
    if current_heading or current_content:
        sections.append(
            Section(
                heading=current_heading,
                level=current_level,
                content="\n\n".join(current_content),
            )
        )

    # If no sections found, put everything in one
    if not sections and current_content:
        sections.append(
            Section(
                heading="Content",
                level=HeadingLevel.LEVEL_1,
                content="\n\n".join(current_content),
            )
        )

    return APADocument(
        title_page=TitlePage(title=title, authors=authors, affiliation=affiliation),
        abstract=abstract_text,
        keywords=keywords,
        sections=sections,
    )


def _extract_formatted_text(paragraph) -> str:
    """Convert paragraph runs to markdown text."""
    result = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue

        # Basic markdown application
        if run.bold:
            text = f"**{text}**"
        if run.italic:
            text = f"*{text}*"

        result.append(text)

    return "".join(result)


def _detect_heading_level(style_name: str, para) -> HeadingLevel | None:
    """Detect if paragraph is a heading, returning the actual level."""
    match = _HEADING_RE.search(style_name)
    if match:
        level_num = int(match.group(1))
        level_map = {
            1: HeadingLevel.LEVEL_1,
            2: HeadingLevel.LEVEL_2,
            3: HeadingLevel.LEVEL_3,
        }
        return level_map.get(level_num, HeadingLevel.LEVEL_1)

    # Heuristic for non-styled documents: short + all bold = level 1
    if "normal" in style_name and len(para.text) < 100:
        runs_with_text = [r for r in para.runs if r.text.strip()]
        if runs_with_text and all(r.bold for r in runs_with_text):
            return HeadingLevel.LEVEL_1

    return None
