"""Word-to-PDF and PDF-to-Word conversion utilities."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document


from apa_formatter.adapters.pdf_adapter import PdfAdapter
from apa_formatter.models.document import APADocument, Section, TitlePage
from apa_formatter.models.enums import (
    DocumentVariant,
    FontChoice,
    HeadingLevel,
    OutputFormat,
)


def docx_to_pdf(source_path: Path, output_path: Path) -> Path:
    """Convert a .docx file to PDF by re-rendering through the APA engine.

    This reads the .docx structure and generates a new APA-formatted PDF.
    Note: only content is converted — original non-APA styling is replaced
    with proper APA 7 formatting during conversion.
    """
    source_path = Path(source_path)
    output_path = Path(output_path)

    if not source_path.exists():
        raise FileNotFoundError(f"Source file not found: {source_path}")

    doc = Document(str(source_path))

    # Extract content from the docx
    title = _extract_title(doc)
    authors = _extract_authors(doc)
    abstract_text = _extract_abstract(doc)
    sections = _extract_sections(doc)

    # Build an APADocument from extracted content
    apa_doc = APADocument(
        title_page=TitlePage(
            title=title,
            authors=authors,
            affiliation="(Converted document)",
            variant=DocumentVariant.STUDENT,
        ),
        abstract=abstract_text if abstract_text else None,
        sections=sections,
        font=FontChoice.TIMES_NEW_ROMAN,
        output_format=OutputFormat.PDF,
    )

    adapter = PdfAdapter(apa_doc)
    return adapter.generate(output_path)


def _extract_title(doc: Any) -> str:
    """Extract the document title from the first non-empty centered paragraph."""
    for para in doc.paragraphs[:10]:
        text = para.text.strip()
        if text and len(text) > 3:
            # Check if the paragraph is bold (likely a title)
            if any(run.bold for run in para.runs if run.text.strip()):
                return str(text)
    # Fallback: use the first non-empty paragraph
    for para in doc.paragraphs:
        if para.text.strip():
            return str(para.text.strip())
    return "Untitled Document"


def _extract_authors(doc: Any) -> list[str]:
    """Extract author names from the title page area."""
    # Look for author-like text after the title
    found_title = False
    for para in doc.paragraphs[:15]:
        text = para.text.strip()
        if not text:
            continue
        if not found_title:
            if any(run.bold for run in para.runs if run.text.strip()):
                found_title = True
                continue
        elif found_title and text and not any(run.bold for run in para.runs if run.text.strip()):
            # Non-bold text after title is likely the author
            return [text]
    return ["Unknown Author"]


def _extract_abstract(doc: Any) -> str | None:
    """Extract abstract text from the document."""
    in_abstract = False
    abstract_parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text.lower() == "abstract":
            in_abstract = True
            continue
        if in_abstract:
            if text and not text.lower().startswith("keyword"):
                # Check if we hit a new section heading (bold centered text)
                is_heading = any(run.bold for run in para.runs if run.text.strip())
                if is_heading and abstract_parts:
                    break
                abstract_parts.append(text)
            elif text.lower().startswith("keyword"):
                break

    return " ".join(abstract_parts) if abstract_parts else None


def _extract_sections(doc: Any) -> list[Section]:
    """Extract body content as sections."""
    sections: list[Section] = []
    current_heading = None
    current_content: list[str] = []

    # Skip title page and abstract
    body_started = False
    past_abstract = False

    for para in doc.paragraphs:
        text = para.text.strip()

        if not body_started:
            if text.lower() == "abstract":
                past_abstract = True
                continue
            if past_abstract and text and not text.lower().startswith("keyword"):
                # Check if this looks like a body heading
                is_bold = any(run.bold for run in para.runs if run.text.strip())
                if is_bold:
                    body_started = True
                    current_heading = text
                    continue
            continue

        if not text:
            continue

        # Check if this paragraph is a heading
        is_bold = any(run.bold for run in para.runs if run.text.strip())

        if is_bold and len(text) < 100:
            # Save previous section
            if current_heading or current_content:
                sections.append(
                    Section(
                        heading=current_heading,
                        level=HeadingLevel.LEVEL_1,
                        content="\n\n".join(current_content),
                    )
                )
            current_heading = text
            current_content = []
        else:
            current_content.append(text)

    # Save last section
    if current_heading or current_content:
        sections.append(
            Section(
                heading=current_heading,
                level=HeadingLevel.LEVEL_1,
                content="\n\n".join(current_content),
            )
        )

    return sections
