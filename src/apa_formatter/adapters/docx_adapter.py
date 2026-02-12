"""Word (.docx) adapter for APA 7 document generation."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from apa_formatter.adapters.base import BaseAdapter
from apa_formatter.domain.models.document import APADocument, Section
from apa_formatter.domain.models.enums import DocumentVariant, HeadingLevel
from apa_formatter.rules.constants import (
    FIRST_LINE_INDENT_INCHES,
    HANGING_INDENT_INCHES,
    HEADING_STYLES,
    LINE_SPACING,
    MARGIN_INCHES,
    REFERENCES_HEADING,
    SPACE_AFTER_PT,
    SPACE_BEFORE_PT,
    TITLE_PAGE_BLANK_LINES_BEFORE_TITLE,
)


class DocxAdapter(BaseAdapter):
    """Generate APA 7 formatted Word documents using python-docx."""

    def __init__(self, document: APADocument, config=None) -> None:
        super().__init__(document, config=config)
        self._docx = Document()
        self._font_spec = self._get_font_spec()

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def generate(self, output_path: Path) -> Path:
        """Build the complete APA 7 document and save to *output_path*."""
        output_path = Path(output_path)
        if not output_path.suffix:
            output_path = output_path.with_suffix(".docx")

        self._setup_page_layout()
        self._setup_default_style()
        self._build_title_page()

        if self.doc.include_toc:
            self._build_table_of_contents()

        if self.doc.abstract:
            self._build_abstract()

        self._build_body()

        if self.doc.references:
            self._build_references()

        if self.doc.appendices:
            self._build_appendices()

        self._docx.save(str(output_path))
        return output_path

    # ------------------------------------------------------------------
    # Table of Contents
    # ------------------------------------------------------------------

    def _build_table_of_contents(self) -> None:
        """Insert a Table of Contents page using Word field codes.

        The TOC is rendered by Word when the document is opened. The user
        will be prompted to update the field on first open.
        """
        # "Table of Contents" heading — centered, bold
        heading_p = self._add_paragraph("")
        heading_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_p.paragraph_format.first_line_indent = Inches(0)
        run = heading_p.add_run("Table of Contents")
        run.bold = True
        run.font.name = self._font_spec.name
        run.font.size = Pt(self._font_spec.size_pt)

        # Insert TOC field code
        toc_p = self._add_paragraph("")
        toc_p.paragraph_format.first_line_indent = Inches(0)
        self._insert_toc_field(toc_p)

        self._add_page_break()

    def _insert_toc_field(self, paragraph) -> None:
        """Insert a TOC field code into the paragraph."""
        run = paragraph.add_run()
        fld_char_begin = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        run._element.append(fld_char_begin)

        instr_run = paragraph.add_run()
        instr_text = instr_run._element.makeelement(qn("w:instrText"), {})
        instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
        instr_run._element.append(instr_text)

        # Separate run to keep field integrity
        sep_run = paragraph.add_run()
        fld_char_sep = sep_run._element.makeelement(
            qn("w:fldChar"), {qn("w:fldCharType"): "separate"}
        )
        sep_run._element.append(fld_char_sep)

        # Placeholder text
        placeholder_run = paragraph.add_run("[Update this field to generate Table of Contents]")
        placeholder_run.font.name = self._font_spec.name
        placeholder_run.font.size = Pt(self._font_spec.size_pt)
        placeholder_run.font.color.rgb = RGBColor(128, 128, 128)

        end_run = paragraph.add_run()
        fld_char_end = end_run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        end_run._element.append(fld_char_end)

    # ------------------------------------------------------------------
    # Page Layout
    # ------------------------------------------------------------------

    def _setup_page_layout(self) -> None:
        """Configure margins, orientation, and page size."""
        section = self._docx.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        section.top_margin = Inches(MARGIN_INCHES)
        section.bottom_margin = Inches(MARGIN_INCHES)
        section.left_margin = Inches(MARGIN_INCHES)
        section.right_margin = Inches(MARGIN_INCHES)

        # Header with page number (top-right)
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Running head for professional papers
        if (
            self.doc.title_page.variant == DocumentVariant.PROFESSIONAL
            and self.doc.title_page.running_head
        ):
            run = hp.add_run(self.doc.title_page.running_head.upper() + "\t")
            run.font.name = self._font_spec.name
            run.font.size = Pt(self._font_spec.size_pt)

        # Page number field
        self._add_page_number_field(hp)

    def _add_page_number_field(self, paragraph) -> None:
        """Insert an auto-updating page number field."""
        run = paragraph.add_run()
        run.font.name = self._font_spec.name
        run.font.size = Pt(self._font_spec.size_pt)

        fld_char_begin = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        run._element.append(fld_char_begin)

        instr_run = paragraph.add_run()
        instr_text = instr_run._element.makeelement(qn("w:instrText"), {})
        instr_text.text = " PAGE "
        instr_run._element.append(instr_text)

        fld_char_end_run = paragraph.add_run()
        fld_char_end = fld_char_end_run._element.makeelement(
            qn("w:fldChar"), {qn("w:fldCharType"): "end"}
        )
        fld_char_end_run._element.append(fld_char_end)

    # ------------------------------------------------------------------
    # Default Style
    # ------------------------------------------------------------------

    def _setup_default_style(self) -> None:
        """Set the default paragraph style for the whole document."""
        style = self._docx.styles["Normal"]
        font = style.font
        font.name = self._font_spec.name
        font.size = Pt(self._font_spec.size_pt)
        font.color.rgb = RGBColor(0, 0, 0)

        pf = style.paragraph_format
        pf.line_spacing = LINE_SPACING
        pf.space_before = Pt(SPACE_BEFORE_PT)
        pf.space_after = Pt(SPACE_AFTER_PT)
        pf.first_line_indent = Inches(FIRST_LINE_INDENT_INCHES)

    # ------------------------------------------------------------------
    # Title Page
    # ------------------------------------------------------------------

    def _build_title_page(self) -> None:
        """Generate the APA 7 title page."""
        tp = self.doc.title_page

        # Blank lines to push title down ~1/3 of page
        for _ in range(TITLE_PAGE_BLANK_LINES_BEFORE_TITLE):
            blank = self._add_paragraph("")
            blank.paragraph_format.first_line_indent = Inches(0)

        # Title — bold, centered, title case
        title_p = self._add_paragraph("")
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.first_line_indent = Inches(0)
        run = title_p.add_run(tp.title)
        run.bold = True
        run.font.name = self._font_spec.name
        run.font.size = Pt(self._font_spec.size_pt)

        # Blank line after title
        self._add_centered_empty()

        # Author name(s)
        authors_text = ", ".join(tp.authors[:-1])
        if len(tp.authors) > 1:
            authors_text += f", and {tp.authors[-1]}"
        else:
            authors_text = tp.authors[0]
        self._add_centered_text(authors_text)

        # Affiliation
        self._add_centered_text(tp.affiliation)

        # Student-specific fields
        if tp.variant == DocumentVariant.STUDENT:
            if tp.course:
                self._add_centered_text(tp.course)
            if tp.instructor:
                self._add_centered_text(tp.instructor)
            if tp.due_date:
                self._add_centered_text(tp.due_date.strftime("%B %d, %Y"))
        else:
            # Professional: author note
            if tp.author_note:
                self._add_centered_empty()
                note_p = self._add_paragraph("")
                note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                note_p.paragraph_format.first_line_indent = Inches(0)
                run = note_p.add_run("Author Note")
                run.bold = True
                run.font.name = self._font_spec.name
                run.font.size = Pt(self._font_spec.size_pt)

                note_body = self._add_paragraph(tp.author_note)
                note_body.paragraph_format.first_line_indent = Inches(FIRST_LINE_INDENT_INCHES)

        # Page break after title page
        self._add_page_break()

    # ------------------------------------------------------------------
    # Abstract
    # ------------------------------------------------------------------

    def _build_abstract(self) -> None:
        """Generate the abstract page."""
        # "Abstract" heading — centered, bold
        heading_p = self._add_paragraph("")
        heading_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_p.paragraph_format.first_line_indent = Inches(0)
        run = heading_p.add_run("Abstract")
        run.bold = True
        run.font.name = self._font_spec.name
        run.font.size = Pt(self._font_spec.size_pt)

        # Abstract text — no indent
        abs_p = self._add_paragraph(self.doc.abstract or "")
        abs_p.paragraph_format.first_line_indent = Inches(0)

        # Keywords (if any)
        if self.doc.keywords:
            kw_p = self._add_paragraph("")
            kw_p.paragraph_format.first_line_indent = Inches(FIRST_LINE_INDENT_INCHES)
            run_label = kw_p.add_run("Keywords: ")
            run_label.italic = True
            run_label.font.name = self._font_spec.name
            run_label.font.size = Pt(self._font_spec.size_pt)
            kw_text = kw_p.add_run(", ".join(self.doc.keywords))
            kw_text.font.name = self._font_spec.name
            kw_text.font.size = Pt(self._font_spec.size_pt)

        self._add_page_break()

    # ------------------------------------------------------------------
    # Body Sections
    # ------------------------------------------------------------------

    def _build_body(self) -> None:
        """Generate all body sections with proper heading levels."""
        for section in self.doc.sections:
            self._render_section(section)

    def _render_section(self, section: Section) -> None:
        """Recursively render a section and its subsections."""
        if section.heading:
            self._add_heading(section.heading, section.level)

        if section.content:
            style = HEADING_STYLES.get(section.level.value)
            # For levels 4-5, content starts on same line as heading (inline)
            # But we already handle that in _add_heading, so just add paragraphs
            if style and style.inline and section.heading:
                # Content was already added inline with heading
                pass
            else:
                # Split content into paragraphs on double newline
                paragraphs = section.content.split("\n\n")
                for para_text in paragraphs:
                    if para_text.strip():
                        self._add_paragraph(para_text.strip())

        for sub in section.subsections:
            self._render_section(sub)

    def _add_heading(self, text: str, level: HeadingLevel) -> None:
        """Add a heading with proper APA 7 formatting for the given level."""
        style = HEADING_STYLES[level.value]
        p = self._add_paragraph("")
        p.paragraph_format.first_line_indent = (
            Inches(FIRST_LINE_INDENT_INCHES) if style.indent else Inches(0)
        )

        if style.centered:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = p.add_run(text)
        run.bold = style.bold
        run.italic = style.italic
        run.font.name = self._font_spec.name
        run.font.size = Pt(self._font_spec.size_pt)

        # For inline headings (levels 4-5), add period and continue text
        if style.inline:
            period_run = p.add_run(". ")
            period_run.bold = False
            period_run.italic = False
            period_run.font.name = self._font_spec.name
            period_run.font.size = Pt(self._font_spec.size_pt)

    # ------------------------------------------------------------------
    # References
    # ------------------------------------------------------------------

    # ------------------------------------------------------------------
    # References
    # ------------------------------------------------------------------

    def _build_references(self) -> None:
        """Generate the APA 7 reference list."""
        self._add_page_break()

        # "References" heading — Level 1 style
        heading_p = self._add_paragraph("")
        heading_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_p.paragraph_format.first_line_indent = Inches(0)
        run = heading_p.add_run(REFERENCES_HEADING)
        run.bold = True
        run.font.name = self._font_spec.name
        run.font.size = Pt(self._font_spec.size_pt)

        # Sort references alphabetically by first author last name
        sorted_refs = sorted(
            self.doc.references,
            key=lambda r: r.authors[0].last_name.lower() if r.authors else "",
        )

        for ref in sorted_refs:
            ref_text = ref.format_apa()
            ref_p = self._add_paragraph("")  # Empty paragraph
            ref_p.paragraph_format.first_line_indent = Inches(-HANGING_INDENT_INCHES)
            ref_p.paragraph_format.left_indent = Inches(HANGING_INDENT_INCHES)

            # Use shared markdown parser
            self._add_markdown_run(ref_p, ref_text)

    def _add_markdown_run(self, paragraph, text: str) -> None:
        """Add text to paragraph, parsing **bold** and *italic* markdown."""
        import re

        # Regex for **bold** or *italic*
        pattern = re.compile(r"(\*\*[^*]+\*\*)|(\*[^*]+\*)")

        last_pos = 0
        for match in pattern.finditer(text):
            # Prefix plain text
            if match.start() > last_pos:
                self._add_run(paragraph, text[last_pos : match.start()])

            content = match.group()
            if content.startswith("**"):
                self._add_run(paragraph, content[2:-2], bold=True)
            else:
                self._add_run(paragraph, content[1:-1], italic=True)

            last_pos = match.end()

        # Suffix
        if last_pos < len(text):
            self._add_run(paragraph, text[last_pos:])

    def _add_run(self, paragraph, text: str, bold: bool = False, italic: bool = False):
        """Helper to add a run with correct font settings."""
        run = paragraph.add_run(text)
        run.font.name = self._font_spec.name
        run.font.size = Pt(self._font_spec.size_pt)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        return run

    # ------------------------------------------------------------------
    # Appendices
    # ------------------------------------------------------------------

    def _build_appendices(self) -> None:
        """Generate appendices."""
        for i, appendix in enumerate(self.doc.appendices):
            self._add_page_break()
            label = f"Appendix {chr(65 + i)}" if len(self.doc.appendices) > 1 else "Appendix"

            heading_p = self._add_paragraph("")
            heading_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading_p.paragraph_format.first_line_indent = Inches(0)
            run = heading_p.add_run(label)
            run.bold = True
            run.font.name = self._font_spec.name
            run.font.size = Pt(self._font_spec.size_pt)

            if appendix.heading:
                title_p = self._add_paragraph("")
                title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_p.paragraph_format.first_line_indent = Inches(0)
                run = title_p.add_run(appendix.heading)
                run.bold = True
                run.font.name = self._font_spec.name
                run.font.size = Pt(self._font_spec.size_pt)

            if appendix.content:
                self._add_paragraph(appendix.content)

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _add_paragraph(self, text: str):
        """Add a paragraph with default APA body style.

        If text contains markdown, it will be parsed.
        """
        p = self._docx.add_paragraph("", style="Normal")
        if text:
            self._add_markdown_run(p, text)
        return p

    def _add_centered_text(self, text: str):
        """Add a centered paragraph (no indent)."""
        p = self._add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        return p

    def _add_centered_empty(self):
        """Add an empty centered paragraph."""
        p = self._add_paragraph("")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        return p

    def _add_page_break(self) -> None:
        """Insert a page break."""
        self._docx.add_page_break()
