"""APA 7 compliance checker for existing .docx files."""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from apa_formatter.rules.constants import (
    FIRST_LINE_INDENT_INCHES,
    FONT_SPECS,
    LINE_SPACING,
    MARGIN_INCHES,
)


@dataclass
class CheckResult:
    """Result of a single APA compliance check."""

    rule: str
    passed: bool
    expected: str
    actual: str
    severity: str = "warning"  # "error" | "warning" | "info"

    @property
    def icon(self) -> str:
        if self.passed:
            return "✅"
        return "❌" if self.severity == "error" else "⚠️"


@dataclass
class ComplianceReport:
    """Full APA 7 compliance report for a document."""

    file_path: str
    results: list[CheckResult] = field(default_factory=list)

    @property
    def total(self) -> int:
        return len(self.results)

    @property
    def passed(self) -> int:
        return sum(1 for r in self.results if r.passed)

    @property
    def failed(self) -> int:
        return self.total - self.passed

    @property
    def score(self) -> float:
        return (self.passed / self.total * 100) if self.total else 0.0

    @property
    def is_compliant(self) -> bool:
        return all(r.passed for r in self.results if r.severity == "error")


class APAChecker:
    """Check an existing .docx file for APA 7 compliance."""

    # Tolerance for comparing measurements (in inches)
    _TOLERANCE = 0.05

    def __init__(self, file_path: Path) -> None:
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"File not found: {self.file_path}")
        if self.file_path.suffix.lower() != ".docx":
            raise ValueError(
                f"Unsupported format: {self.file_path.suffix}. Only .docx is supported."
            )
        self._doc = Document(str(self.file_path))

    def check(self) -> ComplianceReport:
        """Run all APA 7 compliance checks and return a report."""
        report = ComplianceReport(file_path=str(self.file_path))

        report.results.extend(self._check_margins())
        report.results.extend(self._check_font())
        report.results.extend(self._check_line_spacing())
        report.results.extend(self._check_paragraph_indent())
        report.results.extend(self._check_page_size())
        report.results.extend(self._check_page_numbers())

        return report

    # ------------------------------------------------------------------
    # Margin checks
    # ------------------------------------------------------------------

    def _check_margins(self) -> list[CheckResult]:
        results = []
        section = self._doc.sections[0]
        expected = MARGIN_INCHES

        margin_checks = [
            ("Top margin", section.top_margin),
            ("Bottom margin", section.bottom_margin),
            ("Left margin", section.left_margin),
            ("Right margin", section.right_margin),
        ]

        for name, actual_emu in margin_checks:
            actual_inches = actual_emu / 914400 if actual_emu else 0
            passed = abs(actual_inches - expected) < self._TOLERANCE
            results.append(
                CheckResult(
                    rule=name,
                    passed=passed,
                    expected=f'{expected}" (2.54 cm)',
                    actual=f'{actual_inches:.2f}" ({actual_inches * 2.54:.2f} cm)',
                    severity="error",
                )
            )

        return results

    # ------------------------------------------------------------------
    # Font checks
    # ------------------------------------------------------------------

    def _check_font(self) -> list[CheckResult]:
        results = []

        # Check default style font
        style = self._doc.styles["Normal"]
        style_font_name = style.font.name if style.font.name else "Not set"
        style_font_size = style.font.size

        valid_fonts = {spec.name: spec.size_pt for spec in FONT_SPECS.values()}
        font_passed = style_font_name in valid_fonts

        results.append(
            CheckResult(
                rule="Default font family",
                passed=font_passed,
                expected="Times New Roman, Calibri, Arial, or Georgia",
                actual=style_font_name,
                severity="error",
            )
        )

        if style_font_size:
            size_pt = style_font_size.pt
            expected_size = valid_fonts.get(style_font_name, 12)
            size_passed = abs(size_pt - expected_size) < 0.5
            results.append(
                CheckResult(
                    rule="Default font size",
                    passed=size_passed,
                    expected=f"{expected_size}pt",
                    actual=f"{size_pt}pt",
                    severity="error",
                )
            )

        # Sample body paragraphs for font consistency
        body_paras = [
            p
            for p in self._doc.paragraphs
            if p.text.strip() and p.style and p.style.name == "Normal"
        ]
        inconsistent_count = 0
        for p in body_paras[:20]:  # Sample first 20 paragraphs
            for run in p.runs:
                if run.font.name and run.font.name not in valid_fonts:
                    inconsistent_count += 1
                    break

        results.append(
            CheckResult(
                rule="Font consistency in body text",
                passed=inconsistent_count == 0,
                expected="All body paragraphs use APA-approved fonts",
                actual=f"{inconsistent_count} paragraph(s) with non-approved fonts"
                if inconsistent_count
                else "Consistent",
                severity="warning",
            )
        )

        return results

    # ------------------------------------------------------------------
    # Line spacing
    # ------------------------------------------------------------------

    def _check_line_spacing(self) -> list[CheckResult]:
        results = []

        style = self._doc.styles["Normal"]
        pf = style.paragraph_format

        if pf.line_spacing is not None:
            passed = abs(pf.line_spacing - LINE_SPACING) < 0.1
            results.append(
                CheckResult(
                    rule="Line spacing (default style)",
                    passed=passed,
                    expected=f"{LINE_SPACING} (double)",
                    actual=str(pf.line_spacing),
                    severity="error",
                )
            )
        else:
            results.append(
                CheckResult(
                    rule="Line spacing (default style)",
                    passed=False,
                    expected=f"{LINE_SPACING} (double)",
                    actual="Not set (system default)",
                    severity="error",
                )
            )

        # Check space before/after — use `is not None` since Pt(0) is falsy
        space_before = pf.space_before.pt if pf.space_before is not None else None
        space_after = pf.space_after.pt if pf.space_after is not None else None

        results.append(
            CheckResult(
                rule="Space before paragraphs",
                passed=space_before is not None and space_before < 1,
                expected="0pt",
                actual=f"{space_before}pt" if space_before is not None else "Not set",
                severity="warning",
            )
        )

        results.append(
            CheckResult(
                rule="Space after paragraphs",
                passed=space_after is not None and space_after < 1,
                expected="0pt",
                actual=f"{space_after}pt" if space_after is not None else "Not set",
                severity="warning",
            )
        )

        return results

    # ------------------------------------------------------------------
    # Paragraph indent
    # ------------------------------------------------------------------

    def _check_paragraph_indent(self) -> list[CheckResult]:
        results = []

        style = self._doc.styles["Normal"]
        pf = style.paragraph_format

        if pf.first_line_indent:
            indent_inches = pf.first_line_indent / 914400
            expected = FIRST_LINE_INDENT_INCHES
            passed = abs(indent_inches - expected) < self._TOLERANCE
            results.append(
                CheckResult(
                    rule="First-line indent",
                    passed=passed,
                    expected=f'{expected}" (1.27 cm)',
                    actual=f'{indent_inches:.2f}" ({indent_inches * 2.54:.2f} cm)',
                    severity="error",
                )
            )
        else:
            results.append(
                CheckResult(
                    rule="First-line indent",
                    passed=False,
                    expected=f'{FIRST_LINE_INDENT_INCHES}" (1.27 cm)',
                    actual="Not set",
                    severity="error",
                )
            )

        return results

    # ------------------------------------------------------------------
    # Page size
    # ------------------------------------------------------------------

    def _check_page_size(self) -> list[CheckResult]:
        results = []
        section = self._doc.sections[0]

        width_inches = section.page_width / 914400 if section.page_width else 0
        height_inches = section.page_height / 914400 if section.page_height else 0

        width_ok = abs(width_inches - 8.5) < self._TOLERANCE
        height_ok = abs(height_inches - 11.0) < self._TOLERANCE

        results.append(
            CheckResult(
                rule="Paper size",
                passed=width_ok and height_ok,
                expected='8.5" × 11" (Letter)',
                actual=f'{width_inches:.1f}" × {height_inches:.1f}"',
                severity="error",
            )
        )

        return results

    # ------------------------------------------------------------------
    # Page numbers
    # ------------------------------------------------------------------

    def _check_page_numbers(self) -> list[CheckResult]:
        results = []
        section = self._doc.sections[0]
        header = section.header

        header_text = " ".join(p.text for p in header.paragraphs).strip()

        # Check if header has content (page number or running head)
        has_content = bool(header_text) or any(
            "PAGE" in (p.text or "").upper() or len(p.runs) > 0 for p in header.paragraphs
        )

        results.append(
            CheckResult(
                rule="Page numbers in header",
                passed=has_content,
                expected="Page number in top-right corner",
                actual="Header has content" if has_content else "No page number detected",
                severity="warning",
            )
        )

        return results
