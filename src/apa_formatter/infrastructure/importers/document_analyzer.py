"""Document analyzer for pre-import intelligence.

Performs deep analysis of a .docx file before importing into the editor,
extracting metadata, structure, references, formatting info, and APA
compliance issues.
"""

from __future__ import annotations

import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from apa_formatter.domain.models.enums import ReferenceType
from apa_formatter.domain.models.reference import Reference


# ---------------------------------------------------------------------------
# Result Data Classes
# ---------------------------------------------------------------------------


@dataclass
class SectionInfo:
    """Information about a detected document section."""

    heading: str
    level: int
    word_count: int
    paragraph_count: int


@dataclass
class DetectedReference:
    """A reference detected during pre-analysis."""

    raw_text: str
    parsed_reference: Reference | None
    confidence: float  # 0.0 to 1.0
    detection_method: str  # "DOI", "ISBN", "URL", "BibTeX", "heuristic"
    include: bool = True  # user can toggle


@dataclass
class APAIssue:
    """An APA compliance issue detected in the document."""

    severity: str  # "info", "warning", "error"
    message: str
    detail: str = ""


@dataclass
class AnalysisResult:
    """Complete analysis result for a .docx document."""

    # File metadata
    file_name: str = ""
    file_size_kb: float = 0.0
    page_count_estimate: int = 0

    # Title page (best-effort extraction)
    detected_title: str = "Documento sin título"
    detected_authors: list[str] = field(default_factory=lambda: ["Autor desconocido"])
    detected_affiliation: str | None = None
    detected_date: str | None = None

    # Structure
    sections: list[SectionInfo] = field(default_factory=list)
    total_words: int = 0
    total_paragraphs: int = 0
    total_characters: int = 0

    # Formatting
    detected_fonts: list[str] = field(default_factory=list)
    detected_line_spacing: float | None = None
    has_page_numbers: bool = False
    has_headers_footers: bool = False

    # References
    detected_references: list[DetectedReference] = field(default_factory=list)
    raw_reference_paragraphs: list[str] = field(default_factory=list)

    # APA compliance
    apa_issues: list[APAIssue] = field(default_factory=list)

    # Abstract
    detected_abstract: str | None = None
    detected_keywords: list[str] = field(default_factory=list)

    # Success flag
    success: bool = True
    error_message: str = ""


# ---------------------------------------------------------------------------
# Standard APA fonts
# ---------------------------------------------------------------------------

_APA_STANDARD_FONTS = {
    "Times New Roman",
    "Calibri",
    "Arial",
    "Georgia",
    "Computer Modern",
}

_HEADING_PATTERNS = re.compile(r"heading\s*(\d)", re.IGNORECASE)

_REFERENCES_HEADINGS = {"referencias", "references", "bibliografía", "bibliography"}

_ABSTRACT_HEADINGS = {"abstract", "resumen"}


# ---------------------------------------------------------------------------
# Analyzer
# ---------------------------------------------------------------------------


class DocumentAnalyzer:
    """Performs deep pre-import analysis of a .docx file."""

    def analyze(self, path: Path) -> AnalysisResult:
        """Analyze a .docx file and return structured results.

        This is a non-destructive, read-only operation.
        """
        result = AnalysisResult()

        # File-level metadata
        result.file_name = path.name
        try:
            result.file_size_kb = round(os.path.getsize(path) / 1024, 1)
        except OSError:
            result.file_size_kb = 0.0

        # Open document
        try:
            doc = Document(str(path))
        except PackageNotFoundError:
            result.success = False
            result.error_message = f"No se pudo abrir el archivo: {path.name}"
            return result
        except Exception as exc:
            result.success = False
            result.error_message = f"Error al abrir el archivo: {exc}"
            return result

        # Core properties
        self._extract_core_properties(doc, result)

        # Analyze paragraphs
        self._analyze_paragraphs(doc, result)

        # Analyze formatting
        self._analyze_formatting(doc, result)

        # Analyze headers/footers
        self._analyze_headers_footers(doc, result)

        # Detect references
        self._detect_references(result)

        # Run APA compliance checks
        self._check_apa_compliance(result)

        return result

    # -- Core properties --

    def _extract_core_properties(self, doc: Document, result: AnalysisResult) -> None:
        """Extract metadata from document core properties."""
        try:
            props = doc.core_properties
            if props.title and props.title.strip():
                result.detected_title = props.title.strip()
            if props.author and props.author.strip():
                # Split by comma, semicolon or " and "
                raw = props.author.strip()
                authors = re.split(r"[;,]|\band\b", raw)
                result.detected_authors = [a.strip() for a in authors if a.strip()]
            if props.last_modified_by:
                result.detected_affiliation = None  # No direct affiliation in props
            if props.modified:
                result.detected_date = props.modified.strftime("%Y-%m-%d")
            elif props.created:
                result.detected_date = props.created.strftime("%Y-%m-%d")
        except Exception:
            pass  # Core properties may be missing/corrupt

    # -- Paragraph analysis --

    def _analyze_paragraphs(self, doc: Document, result: AnalysisResult) -> None:
        """Walk all paragraphs to extract structure, content, and section info."""
        current_heading: str | None = None
        current_level: int = 1
        current_word_count: int = 0
        current_para_count: int = 0

        total_words = 0
        total_paragraphs = 0
        total_characters = 0

        in_references_section = False
        in_abstract_section = False
        abstract_lines: list[str] = []
        page_break_count = 0
        title_found = False

        for para in doc.paragraphs:
            text = para.text.strip()

            # Count page breaks
            for run in para.runs:
                if (
                    run._element.xml
                    and "w:br" in run._element.xml
                    and 'w:type="page"' in run._element.xml
                ):
                    page_break_count += 1

            if not text:
                continue

            total_paragraphs += 1
            words = len(text.split())
            total_words += words
            total_characters += len(text)

            style_name = (para.style.name if para.style else "").lower()

            # Detect heading
            heading_level = self._get_heading_level(style_name, para)

            if heading_level:
                # Save previous section
                if current_heading is not None:
                    result.sections.append(
                        SectionInfo(
                            heading=current_heading,
                            level=current_level,
                            word_count=current_word_count,
                            paragraph_count=current_para_count,
                        )
                    )

                current_heading = text
                current_level = heading_level
                current_word_count = 0
                current_para_count = 0

                # Check for references section
                in_references_section = text.lower().strip() in _REFERENCES_HEADINGS

                # Check for abstract section
                in_abstract_section = text.lower().strip() in _ABSTRACT_HEADINGS

            elif "title" in style_name and not title_found:
                # Title style detection
                result.detected_title = text
                title_found = True
            else:
                # Regular body text
                current_word_count += words
                current_para_count += 1

                if in_references_section:
                    result.raw_reference_paragraphs.append(text)
                elif in_abstract_section:
                    abstract_lines.append(text)

        # Save last section
        if current_heading is not None:
            result.sections.append(
                SectionInfo(
                    heading=current_heading,
                    level=current_level,
                    word_count=current_word_count,
                    paragraph_count=current_para_count,
                )
            )

        result.total_words = total_words
        result.total_paragraphs = total_paragraphs
        result.total_characters = total_characters
        result.page_count_estimate = max(1, page_break_count + 1)

        # Abstract
        if abstract_lines:
            result.detected_abstract = "\n".join(abstract_lines)
            # Try to detect keywords on last line
            last_line = abstract_lines[-1]
            kw_match = re.match(
                r"(?:keywords?|palabras?\s*clave):\s*(.+)",
                last_line,
                re.IGNORECASE,
            )
            if kw_match:
                kw_text = kw_match.group(1)
                result.detected_keywords = [k.strip() for k in kw_text.split(",") if k.strip()]
                # Remove keyword line from abstract
                result.detected_abstract = "\n".join(abstract_lines[:-1])

        # If title was not found in styles, try first heading or first bold paragraph
        if result.detected_title == "Documento sin título":
            for para in doc.paragraphs[:10]:
                t = para.text.strip()
                if not t:
                    continue
                runs_with_text = [r for r in para.runs if r.text.strip()]
                if runs_with_text and all(r.bold for r in runs_with_text):
                    result.detected_title = t
                    break
                if para.alignment is not None and para.alignment == 1:  # CENTER
                    result.detected_title = t
                    break

    def _get_heading_level(self, style_name: str, para) -> int | None:
        """Detect heading level from paragraph style or formatting."""
        match = _HEADING_PATTERNS.search(style_name)
        if match:
            return int(match.group(1))

        # Heuristic: short + all bold = level 1 heading
        if "normal" in style_name and len(para.text) < 100:
            runs_with_text = [r for r in para.runs if r.text.strip()]
            if runs_with_text and all(r.bold for r in runs_with_text):
                return 1

        return None

    # -- Formatting analysis --

    def _analyze_formatting(self, doc: Document, result: AnalysisResult) -> None:
        """Detect fonts, line spacing, and other formatting details."""
        fonts: set[str] = set()
        spacings: list[float] = []

        for para in doc.paragraphs:
            # Font detection from runs
            for run in para.runs:
                if run.font and run.font.name:
                    fonts.add(run.font.name)

            # Line spacing detection
            fmt = para.paragraph_format
            if fmt and fmt.line_spacing is not None:
                try:
                    spacing_val = float(fmt.line_spacing)
                    if 0.5 < spacing_val < 5.0:
                        spacings.append(spacing_val)
                except (TypeError, ValueError):
                    pass

        result.detected_fonts = sorted(fonts)

        if spacings:
            # Most common spacing
            from collections import Counter

            counter = Counter(spacings)
            result.detected_line_spacing = counter.most_common(1)[0][0]

    def _analyze_headers_footers(self, doc: Document, result: AnalysisResult) -> None:
        """Check for headers and footers."""
        try:
            for section in doc.sections:
                header = section.header
                if header and header.paragraphs:
                    for p in header.paragraphs:
                        if p.text.strip():
                            result.has_headers_footers = True
                            result.has_page_numbers = True
                            break
                footer = section.footer
                if footer and footer.paragraphs:
                    for p in footer.paragraphs:
                        if p.text.strip():
                            result.has_headers_footers = True
                            break
        except Exception:
            pass

    # -- Reference detection --

    def _detect_references(self, result: AnalysisResult) -> None:
        """Parse raw reference paragraphs using SmartReferenceParser."""
        if not result.raw_reference_paragraphs:
            return

        try:
            from apa_formatter.infrastructure.importers.smart_parser import SmartReferenceParser

            parser = SmartReferenceParser()
        except Exception:
            return

        for raw_text in result.raw_reference_paragraphs:
            if not raw_text.strip():
                continue

            confidence = 0.0
            method = "heuristic"
            parsed: Reference | None = None

            try:
                # Check for DOI
                doi_match = parser._DOI_REGEX.search(raw_text)
                if doi_match:
                    method = "DOI"
                    confidence = 0.9
                    try:
                        parsed = parser._doi_fetcher.fetch(doi_match.group(1))
                        if parsed:
                            confidence = 1.0
                    except Exception:
                        pass

                # Check for ISBN
                if not parsed:
                    isbn_match = parser._ISBN_REGEX.search(raw_text)
                    if isbn_match:
                        method = "ISBN"
                        confidence = 0.85
                        try:
                            import re as _re

                            clean_isbn = _re.sub(r"[^0-9X]", "", isbn_match.group(0))
                            parsed = parser._isbn_fetcher.fetch(clean_isbn)
                            if parsed:
                                confidence = 1.0
                        except Exception:
                            pass

                # Try BibTeX
                if not parsed and "@" in raw_text:
                    method = "BibTeX"
                    confidence = 0.7
                    parsed = parser._try_bibtex(raw_text)
                    if parsed:
                        confidence = 0.85

                # Fallback: heuristic
                if not parsed:
                    method = "heuristic"
                    parsed = parser._try_heuristic(raw_text)
                    confidence = 0.4

            except Exception:
                confidence = 0.1

            result.detected_references.append(
                DetectedReference(
                    raw_text=raw_text,
                    parsed_reference=parsed,
                    confidence=confidence,
                    detection_method=method,
                )
            )

    # -- APA compliance checks --

    def _check_apa_compliance(self, result: AnalysisResult) -> None:
        """Run basic APA compliance checks on the analysis results."""
        issues = result.apa_issues

        # Font checks
        non_standard = [f for f in result.detected_fonts if f not in _APA_STANDARD_FONTS]
        if non_standard:
            issues.append(
                APAIssue(
                    severity="warning",
                    message="Fuentes no estándar APA detectadas",
                    detail=f"Se encontraron: {', '.join(non_standard)}. "
                    f"APA 7 recomienda: Times New Roman 12pt, Calibri 11pt, o Arial 11pt.",
                )
            )

        # Line spacing check
        if result.detected_line_spacing and result.detected_line_spacing != 2.0:
            issues.append(
                APAIssue(
                    severity="warning",
                    message="Interlineado diferente a doble espacio",
                    detail=f"Interlineado detectado: {result.detected_line_spacing}. "
                    f"APA 7 requiere interlineado doble (2.0).",
                )
            )

        # Abstract check
        if not result.detected_abstract:
            has_abstract_heading = any(
                s.heading.lower().strip() in _ABSTRACT_HEADINGS for s in result.sections
            )
            if not has_abstract_heading:
                issues.append(
                    APAIssue(
                        severity="info",
                        message="No se detectó abstract/resumen",
                        detail="La mayoría de documentos APA requieren un resumen de 150-250 palabras.",
                    )
                )

        # References check
        if not result.raw_reference_paragraphs:
            issues.append(
                APAIssue(
                    severity="info",
                    message="No se detectó sección de referencias",
                    detail="No se encontró una sección titulada 'Referencias' o 'References'.",
                )
            )

        # Title check
        if result.detected_title == "Documento sin título":
            issues.append(
                APAIssue(
                    severity="warning",
                    message="No se pudo detectar el título del documento",
                    detail="No se encontró un párrafo con estilo 'Title' ni texto centrado en negrita.",
                )
            )

        # Word count info
        if result.total_words > 0:
            if result.total_words < 500:
                issues.append(
                    APAIssue(
                        severity="info",
                        message=f"Documento corto ({result.total_words} palabras)",
                        detail="Los documentos académicos típicamente tienen al menos 1000 palabras.",
                    )
                )

        # Page numbers
        if not result.has_page_numbers:
            issues.append(
                APAIssue(
                    severity="info",
                    message="No se detectaron números de página",
                    detail="APA 7 requiere números de página en la esquina superior derecha.",
                )
            )
